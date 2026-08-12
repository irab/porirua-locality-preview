#!/usr/bin/env python3
"""Build Scope of Work v2 as a .docx for Google Docs / Word upload."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, RGBColor, Cm

OUT = Path(__file__).resolve().parents[1] / "porirua-services-directory-requirements-v2.docx"

INK = RGBColor(0x60, 0x16, 0x4C)
ACCENT = RGBColor(0xCE, 0x20, 0x26)
MUTED = RGBColor(0x8A, 0x67, 0x7A)
TEXT = RGBColor(0x3C, 0x1B, 0x30)
HEADER_BG = "60164C"
ALT_ROW = "FCF1F0"


def set_run_font(run, *, size=11, bold=False, color=TEXT, name="Calibri"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    tcPr.append(shd)


def set_cell_text(cell, text, *, bold=False, color=TEXT, size=10, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=9)
        shade_cell(cell, HEADER_BG)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            # Allow (text, bold) tuples
            if isinstance(val, tuple):
                text, is_bold = val
            else:
                text, is_bold = val, False
            set_cell_text(cell, text, bold=is_bold, size=10)
            if r_idx % 2 == 1:
                shade_cell(cell, ALT_ROW)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True, color=INK, name="Georgia")
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, *, bold=False, size=11, space_after=6, color=TEXT):
    p = doc.add_paragraph()
    # Support simple **bold** segments
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        set_run_font(run, size=size, bold=(bold or i % 2 == 1), color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        parts = item.split("**")
        for i, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            set_run_font(run, size=11, bold=(i % 2 == 1), color=TEXT)
        p.paragraph_format.space_after = Pt(3)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        parts = item.split("**")
        for i, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            set_run_font(run, size=11, bold=(i % 2 == 1), color=TEXT)
        p.paragraph_format.space_after = Pt(3)


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, ALT_ROW)
    cell.text = ""
    p = cell.paragraphs[0]
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        set_run_font(run, size=10, bold=(i % 2 == 1), color=INK)
    doc.add_paragraph()


def build():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    # —— Cover ——
    eyebrow = doc.add_paragraph()
    run = eyebrow.add_run("PORIRUA LOCALITY · TE WĀHI TIAKI TĀTOU")
    set_run_font(run, size=10, bold=True, color=ACCENT)
    eyebrow.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph()
    run = title.add_run("Your Porirua Directory")
    set_run_font(run, size=32, bold=True, color=INK, name="Georgia")
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    run = subtitle.add_run("Scope of Work")
    set_run_font(run, size=18, bold=True, color=INK, name="Georgia")
    subtitle.paragraph_format.space_after = Pt(14)

    add_para(
        doc,
        "A simple, trustworthy online directory so people in Porirua can find help, "
        "connect with community groups, and find local organisations — delivered in two clear phases.",
        size=12,
        space_after=16,
    )

    add_table(
        doc,
        ["", ""],
        [
            [("Prepared for", True), "Porirua Locality team"],
            [("Delivered by", True), "Coshop Limited (NZBN 9429051598860 · Co. 8906928)"],
            [("Version", True), "2.0 · July 2026"],
            [("Timeline", True), "~5–6 weeks (MVP first)"],
            [("Total budget", True), "NZ$10,000"],
            [("Phase 1 — MVP", True), "NZ$5,000 · ~50 hours"],
            [("Phase 2 — Iteration", True), "NZ$5,000 · ~50 hours"],
        ],
        col_widths=[Inches(2.4), Inches(4.0)],
    )

    add_para(
        doc,
        "**Deliverables-focused scope** — what will be built, when, and for how much — without technical detail.",
        size=10,
        color=MUTED,
    )

    doc.add_page_break()

    # —— Contents (Google Docs can replace with Insert → Table of contents) ——
    add_heading(doc, "Contents", level=1)
    add_para(
        doc,
        "This scope sets out what the Porirua Locality team will receive in each phase, "
        "who it is for, and what success looks like.",
        color=MUTED,
        size=10,
    )
    add_para(
        doc,
        "Tip for Google Docs: after upload, place the cursor here and use "
        "Insert → Table of contents to refresh links from the headings below.",
        size=9,
        color=MUTED,
    )

    toc_items = [
        "1. Delivery partner",
        "2. Purpose",
        "3. Why two phases",
        "4. Phase 1 — MVP deliverables",
        "5. Phase 2 — Iteration deliverables",
        "6. Who it is for",
        "7. What success looks like",
        "8. What people will see",
        "9. Data",
        "10. Timeline",
        "11. Budget",
        "12. Not included",
        "13. Assumptions",
        "14. Approval",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        for run in p.runs:
            set_run_font(run, size=11, bold=True, color=INK)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # —— 1 Delivery partner ——
    add_heading(doc, "1. Delivery partner", level=1)
    add_para(
        doc,
        "Work under this scope of work will be performed by **Coshop Limited**, a New Zealand limited company.",
    )
    add_table(
        doc,
        ["Detail", "Information"],
        [
            [("Legal name", True), "Coshop Limited"],
            [("NZBN", True), "9429051598860"],
            [("Companies Office number", True), "8906928"],
            [("Date registered", True), "12 September 2023"],
            [
                ("Registered office", True),
                "Wellington, New Zealand (current address on the Companies Register)",
            ],
            [("Website", True), "coshop.nz"],
        ],
    )
    add_para(
        doc,
        "Coshop builds community-facing digital services — including the CoShop platform for local food hubs — "
        "and will design, build, and hand over **Your Porirua Directory** under the phases and budget below.",
    )

    # —— 2 Purpose ——
    add_heading(doc, "2. Purpose", level=1)
    add_para(
        doc,
        "Build a simple, trustworthy online directory so people in Porirua can "
        "**find help**, **connect with community groups**, and **find local organisations** — "
        "marae, councils, Pātaka Kai, services, and similar — not only urgent crisis help.",
    )
    add_para(doc, "The directory brings together:")
    add_bullets(
        doc,
        [
            "The **Porirua Community Connections Map** — local groups and services the community already knows and trusts",
            "The **New Zealand Family Services Directory** — the government’s national register of social services, filtered to Porirua",
        ],
    )
    add_para(
        doc,
        "It should feel easy on a phone, use plain language, and stay accurate for "
        "**at least three years** with modest ongoing effort from the Porirua Locality team.",
    )
    add_callout(
        doc,
        "Delivery is **two phases**: a testable MVP at half the budget goes live first; "
        "the second half funds changes from real feedback and completes the admin and update layer.",
    )

    # —— 2 ——
    add_heading(doc, "3. Why two phases", level=1)
    add_table(
        doc,
        ["Phase", "Budget", "Goal"],
        [
            [("Phase 1 — MVP", True), "~50 hrs / NZ$5,000", "A working directory people can use and test"],
            [
                ("Phase 2 — Iteration", True),
                "~50 hrs / NZ$5,000",
                "Improve from feedback; add admin, automation, and polish",
            ],
        ],
    )
    add_para(
        doc,
        "Real users — help-seekers and the Locality team — should shape categories, layout, "
        "and content before we invest in the full editing and update setup.",
        size=10,
        color=MUTED,
    )

    # —— 3 ——
    add_heading(doc, "4. Phase 1 — MVP deliverables", level=1)
    add_callout(doc, "~50 hours / NZ$5,000 · ~2–3 weeks to soft launch")
    add_heading(doc, "What people can use", level=2)
    add_bullets(
        doc,
        [
            "**Public website** — shareable URL, works well on phones",
            "**Need-based categories**, **search**, **map**, and **crisis numbers**",
            "**Combined listings** from the Connections Map and the Family Services Directory (Porirua only)",
            "**Clear service cards** — name, short description, phone, address, and website where available",
        ],
    )
    add_heading(doc, "How data is managed in MVP", level=2)
    add_bullets(
        doc,
        [
            "Connections Map stays in the **existing Google Sheet** (the team already knows this)",
            "Government directory imported **once** (can be re-run by hand if needed)",
            "One **published list** the website reads",
            "Basic duplicate flagging; tidy-up by hand is fine for MVP",
        ],
    )

    # —— 4 ——
    add_heading(doc, "5. Phase 2 — Iteration deliverables", level=1)
    add_callout(doc, "~50 hours / NZ$5,000 · ~2–3 weeks after feedback is collected")
    add_bullets(
        doc,
        [
            "Changes from testing (categories, labels, layout, crisis numbers display, map behaviour)",
            "**Password-protected admin** to review, edit, publish, or hide listings",
            "**Weekly automated refresh** of the government directory, with a review queue",
            "**Duplicate linking** in admin (so the same service is not shown twice)",
            "**Embed** on the Porirua Locality website",
            "Short **editor guide**, backups, and handover so the directory lasts",
        ],
    )

    # —— 5 ——
    add_heading(doc, "6. Who it is for", level=1)
    add_heading(doc, "Public (no login, no personal details)", level=2)
    add_table(
        doc,
        ["Group", "Need"],
        [
            (
                ("Immediate help", True),
                "Support for themselves or someone they know (food, housing, safety, health, and similar)",
            ),
            (
                ("Community connection", True),
                "Contact or connect with community groups and local kaupapa",
            ),
            (
                ("Civic & community places", True),
                "Marae, councils, Pātaka Kai, libraries, and organisations the community trusts",
            ),
        ],
    )
    add_heading(doc, "Behind the scenes", level=2)
    add_para(
        doc,
        "Porirua Locality team members keep listings accurate — Sheet and manual updates in Phase 1; "
        "admin tools in Phase 2.",
    )
    add_para(
        doc,
        "Community workers may share the link. This is **not** a referral or case-management system.",
    )

    # —— 6 ——
    add_heading(doc, "7. What success looks like", level=1)
    add_table(
        doc,
        ["Goal", "How we will know"],
        [
            ("People find help quickly", "Clear categories, search, and a map of nearby services"),
            ("Information is trustworthy", "Curated merge for MVP; publish/hide workflow in Phase 2"),
            (
                "Local context is kept",
                "Connections Map organisations keep their local story and kaupapa where relevant",
            ),
            ("Real-world validation", "MVP live within ~50 hours; feedback collected before Phase 2"),
            (
                "Low ongoing effort",
                "After Phase 2: automatic government refresh; team ~1–2 hours per month on review",
            ),
            (
                "Long-lasting",
                "Stable tools, exportable data, documented handover — useful for ~3 years",
            ),
            ("Fits the budget", "MVP ≤ NZ$5,000; total ≤ NZ$10,000"),
        ],
    )

    # —— 7 ——
    add_heading(doc, "8. What people will see", level=1)
    add_para(doc, "When someone opens the directory they can:")
    add_numbered(
        doc,
        [
            "**Choose a path** — find support, or connect with community",
            "**Browse by need** — plain-language categories such as food / kai, housing, money help, feeling unsafe, counselling, health, legal advice, work and learning, everyday needs",
            "**Search** — type what they need",
            "**Use the map** — see nearby services when locations are known",
            "**Reach crisis numbers** — always easy to find",
            "**Open a listing** — name, short description, address or suburb, phone, website, and opening hours if known",
            "**Save a short list for this visit** (optional) — nothing stored online; print if useful",
        ],
    )
    add_para(doc, "People should not need to know which list a service came from.")
    add_table(
        doc,
        ["", "Phase 1", "Phase 2"],
        [
            [("Standalone website", True), "Yes", "Yes"],
            [("On the Porirua Locality site", True), "Link out", "Embed"],
        ],
    )

    # —— 8 ——
    add_heading(doc, "9. Data", level=1)
    add_table(
        doc,
        ["Source", "What it provides"],
        [
            (("Community Connections Map", True), "Local organisations, kaupapa, and initiatives"),
            (("NZ Family Services Directory", True), "Broader social services, filtered to Porirua"),
        ],
    )
    add_para(
        doc,
        "Only services relevant to **people in Porirua** appear. When the same organisation is in both lists, "
        "we prefer the **richest local description** (usually from the Connections Map) and avoid showing the same service twice.",
    )

    # —— 9 ——
    add_heading(doc, "10. Timeline", level=1)
    add_heading(doc, "Phase 1 — MVP (~2–3 weeks)", level=2)
    add_table(
        doc,
        ["Week", "Focus"],
        [
            (("1", True), "Bring the two data sources together; categories and crisis numbers"),
            (("2", True), "Search, map, service cards; mobile check; public URL live"),
            (("3", True), "Soft launch, fix gaps, brief handover — start user testing"),
        ],
    )
    add_callout(
        doc,
        "**Between phases (~1–2 weeks, not in budget)** — Collect feedback from help-seekers, "
        "the Locality team, and community workers. Prioritise changes for Phase 2.",
    )
    add_heading(doc, "Phase 2 — Iteration (~2–3 weeks)", level=2)
    add_table(
        doc,
        ["Week", "Focus"],
        [
            (("1", True), "Priority changes from feedback; admin for review and publish"),
            (("2", True), "Weekly government refresh; duplicates; Locality site embed"),
            (("3", True), "Editor guide, backups, production launch"),
        ],
    )
    add_para(doc, "**Total elapsed:** roughly 5–8 weeks including the feedback window.")

    # —— 10 ——
    add_heading(doc, "11. Budget", level=1)
    add_table(
        doc,
        ["Phase", "Hours", "Cost"],
        [
            [("Phase 1 — MVP", True), "~50", ("NZ$5,000", True)],
            [("Phase 2 — Iteration", True), "~50", ("NZ$5,000", True)],
            [("Total", True), "~100", ("NZ$10,000", True)],
        ],
    )
    add_para(
        doc,
        "**Ongoing hosting (from Phase 2):** about NZ$25–60 per month. "
        "MVP can run at minimal cost on a simple public website.",
    )

    # —— 11 ——
    add_heading(doc, "12. Not included", level=1)
    add_bullets(
        doc,
        [
            "Phone helpline or human navigator",
            "Referral tracking (“did the person get help?”)",
            "Public user accounts or collecting personal information",
            "Provider login to edit their own listing",
            "Installable phone app",
            "Usage analytics dashboard",
            "Te Reo category labels (can be added later)",
            "Extra data sources beyond Connections Map and Family Services Directory",
            "National coverage outside Porirua",
            "Replacing 111 or national crisis helplines",
            "Real-time availability (e.g. “bed free tonight”)",
        ],
    )
    add_callout(
        doc,
        "**Phase 1 specifically does not include:** admin website, automatic weekly refresh, "
        "Locality site embed, formal duplicate tools, or full editor runbooks — those are Phase 2.",
    )

    # —— 12 ——
    add_heading(doc, "13. Assumptions", level=1)
    add_table(
        doc,
        ["Item", "Assumption"],
        [
            ("Porirua Locality team", "People available to review listings; ~1–2 hours per month after Phase 2"),
            ("Government directory", "Continues to be published as open data"),
            (
                "Connections Map",
                "Existing Google Sheet remains the starting point for local organisations",
            ),
            ("Locality website", "Can host an embed in Phase 2"),
            ("Branding & content", "Team provides crisis numbers and final category wording"),
            ("MVP feedback", "Team helps recruit testers and capture feedback between phases"),
        ],
    )

    # —— 14 ——
    add_heading(doc, "14. Approval", level=1)
    add_table(
        doc,
        ["Role", "Name", "Date", "Signature"],
        [
            ("Project sponsor", "", "", ""),
            ("Porirua Locality lead", "", "", ""),
            ("Contractor — Coshop Limited", "", "", ""),
        ],
    )

    footer = doc.add_paragraph()
    run = footer.add_run(
        "Your Porirua Directory — Scope of Work v2.0 · July 2026 · Porirua Locality / Te Wāhi Tiaki Tātou"
    )
    set_run_font(run, size=9, color=MUTED)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
